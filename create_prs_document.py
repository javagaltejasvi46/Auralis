#!/usr/bin/env python3
"""
Script to create Auralis PRS Document in DOCX format
This script generates a comprehensive Product Requirements Specification
document containing all features from the documentation folder.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def create_heading(doc, text, level=1):
    """Create a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def create_table_of_contents(doc):
    """Create a table of contents"""
    create_heading(doc, "TABLE OF CONTENTS", 1)
    
    toc_items = [
        ("1. Executive Summary", "3"),
        ("2. System Overview", "5"),
        ("3. Feature Specifications", "7"),
        ("   3.1 Authentication System", "8"),
        ("   3.2 Enhanced Patient Report System", "12"),
        ("   3.3 AI Notes Generation", "16"),
        ("   3.4 Patient Search", "20"),
        ("   3.5 Session Management", "23"),
        ("   3.6 Real-Time Transcription", "27"),
        ("   3.7 Phi-3-Mini Summarization Migration", "31"),
        ("4. Technical Architecture", "35"),
        ("5. Implementation Status", "38"),
        ("6. Quality Assurance", "41"),
        ("7. Deployment Requirements", "44"),
        ("8. Appendices", "47")
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).bold = False
        p.add_run(f" {'.' * (60 - len(item))} {page}")
    
    add_page_break(doc)

def create_prs_document():
    """Create the complete PRS document"""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading("AURALIS MEDICAL VOICE TRANSCRIPTION SYSTEM", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("PRODUCT REQUIREMENTS SPECIFICATION (PRS)", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Document metadata
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run("Document Version: 1.0\n").bold = True
    metadata.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n").bold = True
    metadata.add_run("Prepared by: Kiro AI Assistant").bold = True
    
    add_page_break(doc)
    
    # Table of Contents
    create_table_of_contents(doc)
    
    # 1. Executive Summary
    create_heading(doc, "1. EXECUTIVE SUMMARY", 1)
    
    doc.add_paragraph(
        "The Auralis Medical Voice Transcription System is a comprehensive HIPAA-compliant "
        "platform designed for mental health professionals to manage therapy sessions, patient "
        "records, and clinical documentation. The system provides real-time transcription, "
        "AI-powered clinical note generation, and comprehensive patient management capabilities."
    )
    
    create_heading(doc, "Key Features", 2)
    features = [
        "Secure Authentication: JWT-based therapist authentication with bcrypt password hashing",
        "Comprehensive Patient Management: 45+ field patient profiles following psychotherapy report standards",
        "AI-Powered Clinical Notes: Automated note generation using locally-hosted Phi-3-Mini model",
        "Intelligent Patient Search: Multi-field search with fuzzy matching across names, IDs, and phone numbers",
        "Complete Session Management: Audio recording, transcription storage, and session lifecycle management",
        "Real-Time Transcription: Live speech-to-text with 90+ language support and speaker diarization",
        "Local AI Processing: Privacy-focused local model deployment eliminating external API dependencies"
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    create_heading(doc, "Business Value", 2)
    values = [
        "Efficiency: Reduces documentation time by 70% through AI automation",
        "Compliance: HIPAA-compliant architecture with local data processing",
        "Scalability: Supports multiple therapists with complete data isolation",
        "Quality: Professional-grade clinical documentation following industry standards",
        "Privacy: All AI processing done locally, no patient data leaves the system"
    ]
    
    for value in values:
        p = doc.add_paragraph(value, style='List Bullet')
    
    add_page_break(doc)
    
    # 2. System Overview
    create_heading(doc, "2. SYSTEM OVERVIEW", 1)
    
    create_heading(doc, "2.1 Architecture Overview", 2)
    doc.add_paragraph(
        "The Auralis system follows a modern microservices architecture with the following components:"
    )
    
    # Architecture layers
    layers = [
        ("Frontend Layer", "React Native Mobile App with TypeScript"),
        ("API Gateway", "FastAPI Backend with JWT Authentication"),
        ("Core Services", "Patient Management, Session Management, Auth Service"),
        ("AI Services", "Phi-3-Mini Summarizer, Faster-Whisper, Search Engine"),
        ("Data Layer", "SQLite Database + File Storage")
    ]
    
    for layer, description in layers:
        p = doc.add_paragraph()
        p.add_run(f"{layer}: ").bold = True
        p.add_run(description)
    
    create_heading(doc, "2.2 Technology Stack", 2)
    
    tech_sections = [
        ("Frontend", [
            "React Native with TypeScript",
            "Expo framework for cross-platform development", 
            "WebSocket client for real-time transcription",
            "Secure token storage"
        ]),
        ("Backend", [
            "FastAPI (Python) for REST API",
            "WebSocket server for real-time communication",
            "SQLite database for data persistence",
            "JWT authentication with bcrypt password hashing"
        ]),
        ("AI/ML Components", [
            "Phi-3-Mini (3.8B parameters) for clinical note generation",
            "Faster-Whisper for speech-to-text transcription",
            "llama-cpp-python for optimized inference",
            "Local model deployment (no external APIs)"
        ]),
        ("Infrastructure", [
            "Docker containerization",
            "Volume mounts for model storage",
            "HIPAA-compliant local deployment"
        ])
    ]
    
    for section, items in tech_sections:
        create_heading(doc, section, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # 3. Feature Specifications
    create_heading(doc, "3. FEATURE SPECIFICATIONS", 1)
    
    # 3.1 Authentication System
    create_heading(doc, "3.1 Authentication System", 2)
    
    create_heading(doc, "3.1.1 Overview", 3)
    doc.add_paragraph(
        "Secure JWT-based authentication system providing therapist registration, login, "
        "and session management with HIPAA-compliant security measures."
    )
    
    create_heading(doc, "3.1.2 Core Requirements", 3)
    
    auth_requirements = [
        ("Therapist Registration", [
            "Email, username, password, and full name (required)",
            "Optional fields: license number, specialization, phone number",
            "Password hashing using bcrypt with cost factor 12",
            "Unique constraints on email and username",
            "Account activation and verification support"
        ]),
        ("Secure Login", [
            "Credential validation with constant-time comparison",
            "JWT token generation with 24-hour expiration",
            "Last login timestamp tracking",
            "Secure error handling (no information leakage)"
        ]),
        ("Token-Based Authorization", [
            "HS256 algorithm with configurable secret key",
            "Therapist ID embedded in token payload",
            "Automatic token validation on protected endpoints",
            "Graceful handling of expired/invalid tokens"
        ]),
        ("Data Isolation", [
            "Complete separation of therapist data",
            "Patient and session filtering by therapist ID",
            "404 responses for unauthorized access (prevents information leakage)",
            "Audit trail for all authentication events"
        ])
    ]
    
    for req_title, req_items in auth_requirements:
        create_heading(doc, req_title, 4)
        for item in req_items:
            doc.add_paragraph(item, style='List Bullet')
    
    create_heading(doc, "3.1.3 Implementation Status", 3)
    doc.add_paragraph("✅ COMPLETED - All authentication features implemented and tested")
    
    completed_items = [
        "Database model with all required fields",
        "Password security with bcrypt implementation", 
        "JWT token management with proper expiration",
        "Complete frontend authentication flow",
        "Data isolation across all endpoints",
        "Comprehensive error handling"
    ]
    
    for item in completed_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 3.2 Enhanced Patient Report System
    create_heading(doc, "3.2 Enhanced Patient Report System", 2)
    
    create_heading(doc, "3.2.1 Overview", 3)
    doc.add_paragraph(
        "Comprehensive patient information management system following professional "
        "psychotherapy report standards with 45+ data fields and PDF export capabilities."
    )
    
    create_heading(doc, "3.2.2 Core Requirements", 3)
    
    patient_sections = [
        ("Core Demographics (8 fields)", [
            "Full name, age, gender, date of birth",
            "Residence, education, occupation, marital status"
        ]),
        ("Medical Information (10 fields)", [
            "Current medical conditions, past medical conditions",
            "Current medications, allergies, hospitalizations",
            "Previous psychiatric diagnoses and treatments",
            "Suicide/self-harm history, substance use history"
        ]),
        ("Family & Social History (12 fields)", [
            "Family psychiatric/medical illness, family dynamics",
            "Childhood development, educational history",
            "Occupational history, relationship history",
            "Social support system, living situation",
            "Cultural/religious background"
        ]),
        ("Clinical Assessment (15 fields)", [
            "Chief complaint and description",
            "Illness onset, progression, previous episodes",
            "Triggers, functional impact",
            "Complete mental status examination (11 components)"
        ])
    ]
    
    for section, items in patient_sections:
        create_heading(doc, section, 4)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    create_heading(doc, "3.2.3 PDF Export Functionality", 3)
    pdf_features = [
        "Complete patient reports in PDF format",
        "Professional psychotherapy report template structure",
        "All patient information sections included",
        "Session summaries with dates and findings",
        "Therapist name and generation date in header",
        "Shareable format for healthcare provider collaboration"
    ]
    
    for feature in pdf_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    create_heading(doc, "3.2.4 Implementation Status", 3)
    doc.add_paragraph("✅ COMPLETED - All patient report features implemented")
    
    # 3.3 AI Notes Generation
    create_heading(doc, "3.3 AI Notes Generation", 2)
    
    create_heading(doc, "3.3.1 Overview", 3)
    doc.add_paragraph(
        "Automated clinical note generation system using locally-hosted Phi-3-Mini "
        "language model to create structured, professional clinical summaries from "
        "therapy session transcriptions."
    )
    
    create_heading(doc, "3.3.2 Clinical Note Structure", 3)
    note_sections = [
        "Chief Complaint: Primary presenting issue",
        "Emotional State: Current mood and affect",
        "Risk Assessment: Safety concerns with highlighted keywords",
        "Intervention: Therapeutic techniques used",
        "Progress: Client advancement and response",
        "Plan: Next session objectives and homework"
    ]
    
    for section in note_sections:
        doc.add_paragraph(section, style='List Bullet')
    
    create_heading(doc, "3.3.3 AI Model Specifications", 3)
    model_specs = [
        "Model: Phi-3-Mini (3.8B parameters)",
        "Deployment: Local Ollama server (port 11434)",
        "Performance: 10-15 seconds per note (CPU), 2-4 seconds (GPU)",
        "Context Window: 4K tokens (sufficient for therapy sessions)",
        "Output Limit: 50 words per section, 150 words total",
        "Temperature: 0.7 (balanced creativity/consistency)"
    ]
    
    for spec in model_specs:
        doc.add_paragraph(spec, style='List Bullet')
    
    create_heading(doc, "3.3.4 Implementation Status", 3)
    doc.add_paragraph("✅ COMPLETED - All AI notes features implemented")
    
    add_page_break(doc)
    
    # 3.4 Patient Search
    create_heading(doc, "3.4 Patient Search", 2)
    
    create_heading(doc, "3.4.1 Overview", 3)
    doc.add_paragraph(
        "Intelligent patient search system providing unified search across multiple "
        "patient identifiers with fuzzy matching, real-time results, and relevance scoring."
    )
    
    create_heading(doc, "3.4.2 Search Capabilities", 3)
    search_features = [
        "Patient Name: Fuzzy matching with typo tolerance",
        "Phone Number: Normalized matching across formats", 
        "Patient ID: Exact 6-digit alphanumeric matching",
        "Mixed Queries: Intelligent type detection",
        "Real-time Results: Search as user types",
        "Relevance Scoring: Prioritized exact matches"
    ]
    
    for feature in search_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    create_heading(doc, "3.4.3 Implementation Status", 3)
    doc.add_paragraph("✅ COMPLETED - All patient search features implemented")
    
    # 3.5 Session Management
    create_heading(doc, "3.5 Session Management", 2)
    
    create_heading(doc, "3.5.1 Overview", 3)
    doc.add_paragraph(
        "Comprehensive therapy session lifecycle management system handling session "
        "creation, audio file storage, transcription management, and clinical documentation."
    )
    
    create_heading(doc, "3.5.2 Core Capabilities", 3)
    session_features = [
        "Auto-incrementing session numbers per patient",
        "Session metadata: date, time, duration, language",
        "Audio file support: WAV, M4A, MP3 (up to 500MB)",
        "Transcription storage (original and translated)",
        "Clinical notes and treatment plans",
        "AI generation metadata tracking",
        "Complete session lifecycle management"
    ]
    
    for feature in session_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    create_heading(doc, "3.5.3 Implementation Status", 3)
    doc.add_paragraph("✅ COMPLETED - All session management features implemented")
    
    # 3.6 Real-Time Transcription
    create_heading(doc, "3.6 Real-Time Transcription", 2)
    
    create_heading(doc, "3.6.1 Overview", 3)
    doc.add_paragraph(
        "Live speech-to-text transcription system using Faster-Whisper with support "
        "for 90+ languages, automatic translation, speaker diarization, and WebSocket-based "
        "real-time streaming."
    )
    
    create_heading(doc, "3.6.2 Key Features", 3)
    transcription_features = [
        "Real-time audio processing via WebSocket (port 8003)",
        "90+ language support with automatic detection",
        "Automatic English translation for non-English audio",
        "Speaker diarization with person labeling",
        "Audio chunk processing (2-3 second intervals)",
        "Support for WAV, M4A, MP3 file formats",
        "High accuracy for Indian languages (Hindi, Tamil, Telugu, Kannada)"
    ]
    
    for feature in transcription_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    create_heading(doc, "3.6.3 Performance Specifications", 3)
    perf_specs = [
        "CPU Processing: 0.5-1.0x real-time speed",
        "GPU Processing: 5-10x real-time speed", 
        "Transcription Accuracy: 90%+ for clear speech",
        "Latency: 2-3 seconds for real-time updates",
        "Language Detection: 95%+ accuracy",
        "Speaker Diarization: 85%+ accuracy for 2-3 speakers"
    ]
    
    for spec in perf_specs:
        doc.add_paragraph(spec, style='List Bullet')
    
    create_heading(doc, "3.6.4 Implementation Status", 3)
    doc.add_paragraph("✅ COMPLETED - All real-time transcription features implemented")
    
    # 3.7 Phi-3-Mini Summarization Migration
    create_heading(doc, "3.7 Phi-3-Mini Summarization Migration", 2)
    
    create_heading(doc, "3.7.1 Overview", 3)
    doc.add_paragraph(
        "Migration from Google Gemini API to locally-hosted Phi-3-Mini model for "
        "therapy session summarization, providing enhanced privacy, reliability, and cost-effectiveness."
    )
    
    create_heading(doc, "3.7.2 Migration Benefits", 3)
    migration_benefits = [
        "Enhanced Privacy: All AI processing done locally",
        "Cost Reduction: Eliminates external API costs",
        "Improved Reliability: No dependency on external services",
        "HIPAA Compliance: Patient data never leaves the system",
        "Performance Optimization: Faster inference with local deployment",
        "Customization: Fine-tuned model for psychotherapy domain"
    ]
    
    for benefit in migration_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    create_heading(doc, "3.7.3 Model Specifications", 3)
    phi3_specs = [
        "Model: Phi-3-Mini-4K-Instruct (3.8B parameters)",
        "Size: ~2.5GB (Q4_K_M quantized)",
        "Training: LoRA fine-tuning on psychotherapy dataset",
        "Inference Engine: llama-cpp-python",
        "Performance: <15 seconds (CPU), <5 seconds (GPU)",
        "Memory Usage: <4GB for inference"
    ]
    
    for spec in phi3_specs:
        doc.add_paragraph(spec, style='List Bullet')
    
    create_heading(doc, "3.7.4 Implementation Status", 3)
    doc.add_paragraph("🔄 IN PROGRESS (75% Complete) - Core infrastructure completed, training pipeline in development")
    
    progress_items = [
        "✅ Model evaluation and selection",
        "✅ Infrastructure setup and configuration",
        "✅ Database schema updates",
        "✅ API endpoint design",
        "🔄 Training pipeline implementation",
        "🔄 Model fine-tuning on dataset",
        "🔄 GGUF conversion and quantization",
        "⏳ Integration testing and deployment"
    ]
    
    for item in progress_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # 4. Technical Architecture
    create_heading(doc, "4. TECHNICAL ARCHITECTURE", 1)
    
    create_heading(doc, "4.1 System Architecture Overview", 2)
    doc.add_paragraph(
        "The Auralis system employs a modern, scalable architecture designed for "
        "healthcare applications with strict privacy and security requirements."
    )
    
    create_heading(doc, "4.1.1 Architecture Principles", 3)
    principles = [
        "Microservices Design: Modular services for scalability and maintainability",
        "Local-First Processing: All AI/ML operations performed locally",
        "HIPAA Compliance: End-to-end encryption and audit trails",
        "Data Isolation: Complete separation between therapist accounts",
        "Fault Tolerance: Graceful degradation and error recovery",
        "Performance Optimization: Sub-second response times for critical operations"
    ]
    
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    create_heading(doc, "4.2 Security Architecture", 2)
    
    security_sections = [
        ("Authentication & Authorization", [
            "JWT Tokens: Stateless authentication with 24-hour expiration",
            "bcrypt Hashing: Password security with cost factor 12",
            "Role-Based Access: Therapist-level data isolation",
            "Session Management: Secure token lifecycle management"
        ]),
        ("Data Protection", [
            "Encryption: AES-256 for data at rest",
            "HTTPS/TLS: All communications encrypted in transit",
            "Local Processing: Patient data never leaves the system",
            "Audit Trails: Comprehensive logging for compliance"
        ]),
        ("HIPAA Compliance", [
            "Administrative Safeguards: Access controls and training",
            "Physical Safeguards: Secure deployment environments",
            "Technical Safeguards: Encryption, audit logs, access controls",
            "Business Associate Agreements: Vendor compliance requirements"
        ])
    ]
    
    for section, items in security_sections:
        create_heading(doc, section, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    create_heading(doc, "4.3 Performance Architecture", 2)
    
    perf_targets = [
        "Authentication: <500ms for login/token validation",
        "Patient Search: <2s for 1000+ patient database",
        "AI Note Generation: <15s (CPU), <5s (GPU)",
        "Real-time Transcription: 2-3s latency",
        "File Upload: Progress tracking for large audio files"
    ]
    
    create_heading(doc, "Response Time Targets", 3)
    for target in perf_targets:
        doc.add_paragraph(target, style='List Bullet')
    
    # 5. Implementation Status
    create_heading(doc, "5. IMPLEMENTATION STATUS", 1)
    
    create_heading(doc, "5.1 Completed Features ✅", 2)
    
    completed_features = [
        ("Authentication System (100% Complete)", [
            "Therapist registration and login",
            "JWT token management with 24-hour expiration",
            "bcrypt password hashing with salt",
            "Data isolation across all endpoints",
            "Frontend authentication flow",
            "Secure token storage and management"
        ]),
        ("Enhanced Patient Report System (100% Complete)", [
            "45+ field patient data model",
            "Multi-section patient registration form",
            "Comprehensive patient profile display",
            "Patient information editing",
            "Overall summary generation",
            "PDF export with professional formatting"
        ]),
        ("AI Notes Generation (100% Complete)", [
            "Phi-3-Mini integration via Ollama",
            "Structured clinical note generation",
            "Risk keyword detection and highlighting",
            "Note editing with markdown preservation",
            "Regeneration functionality",
            "Complete metadata tracking"
        ]),
        ("Patient Search (100% Complete)", [
            "Multi-field search (name, phone, ID)",
            "Fuzzy matching with typo tolerance",
            "Real-time search results",
            "Relevance scoring and prioritization",
            "Search result highlighting"
        ]),
        ("Session Management (100% Complete)", [
            "Complete session lifecycle management",
            "Audio file upload and storage",
            "Auto-incrementing session numbers",
            "Transcription and translation storage",
            "AI metadata tracking",
            "Session deletion with file cleanup"
        ]),
        ("Real-Time Transcription (100% Complete)", [
            "Faster-Whisper integration",
            "WebSocket server for real-time streaming",
            "90+ language support with auto-detection",
            "Automatic English translation",
            "Speaker diarization with person labeling",
            "Audio file processing support"
        ])
    ]
    
    for feature, items in completed_features:
        create_heading(doc, feature, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    create_heading(doc, "5.2 In Progress Features 🔄", 2)
    
    create_heading(doc, "Phi-3-Mini Summarization Migration (75% Complete)", 3)
    progress_status = [
        "✅ Model evaluation and selection",
        "✅ Infrastructure setup and configuration", 
        "✅ Database schema updates",
        "✅ API endpoint design",
        "🔄 Training pipeline implementation",
        "🔄 Model fine-tuning on dataset",
        "🔄 GGUF conversion and quantization",
        "⏳ Integration testing and deployment"
    ]
    
    for status in progress_status:
        doc.add_paragraph(status, style='List Bullet')
    
    create_heading(doc, "5.3 Overall System Completion: 96%", 2)
    
    # Create completion table
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Completion'
    hdr_cells[3].text = 'Key Capabilities'
    
    # Data rows
    completion_data = [
        ('Authentication System', '✅ Complete', '100%', 'JWT auth, bcrypt, data isolation'),
        ('Patient Report System', '✅ Complete', '100%', '45+ fields, PDF export, summaries'),
        ('AI Notes Generation', '✅ Complete', '100%', 'Phi-3 notes, risk detection, editing'),
        ('Patient Search', '✅ Complete', '100%', 'Multi-field, fuzzy matching, real-time'),
        ('Session Management', '✅ Complete', '100%', 'Full lifecycle, audio, metadata'),
        ('Real-Time Transcription', '✅ Complete', '100%', '90+ languages, WebSocket, diarization'),
        ('Phi-3 Migration', '🔄 In Progress', '75%', 'Local model, training pipeline')
    ]
    
    for i, (feature, status, completion, capabilities) in enumerate(completion_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = feature
        row_cells[1].text = status
        row_cells[2].text = completion
        row_cells[3].text = capabilities
    
    add_page_break(doc)
    
    # 6. Quality Assurance
    create_heading(doc, "6. QUALITY ASSURANCE", 1)
    
    create_heading(doc, "6.1 Testing Strategy", 2)
    doc.add_paragraph(
        "The Auralis system employs a comprehensive testing strategy ensuring "
        "reliability, security, and performance across all components."
    )
    
    create_heading(doc, "6.1.1 Testing Pyramid", 3)
    
    testing_levels = [
        ("Unit Testing (Foundation)", [
            "Individual component functionality",
            "Business logic validation", 
            "Error handling verification",
            "Mock external dependencies",
            "Target: 90%+ code coverage"
        ]),
        ("Property-Based Testing (Robustness)", [
            "Universal properties across all inputs",
            "Hypothesis (Python) and fast-check (TypeScript)",
            "100+ iterations per property test",
            "Edge case discovery and validation",
            "Correctness guarantee verification"
        ]),
        ("Integration Testing (System)", [
            "End-to-end workflow validation",
            "API endpoint integration",
            "Database transaction integrity",
            "File system operations",
            "Cross-service communication"
        ]),
        ("Performance Testing (Scalability)", [
            "Response time benchmarking",
            "Concurrent user simulation",
            "Memory usage profiling",
            "AI model inference timing",
            "Database query optimization"
        ])
    ]
    
    for level, items in testing_levels:
        create_heading(doc, level, 4)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    create_heading(doc, "6.2 Test Coverage by Feature", 2)
    
    test_coverage = [
        ("Authentication System", "17 comprehensive test cases"),
        ("Enhanced Patient Report", "15 test cases covering 45+ fields"),
        ("AI Notes Generation", "14 test cases for clinical note generation"),
        ("Patient Search", "13 test cases for intelligent search"),
        ("Session Management", "16 test cases for session lifecycle"),
        ("Real-Time Transcription", "12 test cases for WebSocket transcription"),
        ("Phi-3 Migration", "11 test cases for model migration")
    ]
    
    for feature, coverage in test_coverage:
        p = doc.add_paragraph()
        p.add_run(f"{feature}: ").bold = True
        p.add_run(coverage)
    
    create_heading(doc, "6.3 Quality Metrics", 2)
    
    quality_sections = [
        ("Performance Benchmarks", [
            "Response Times: 50ms-2s depending on operation",
            "Concurrent Users: 5-20+ depending on hardware",
            "AI Generation: 10-15s (CPU), 2-4s (GPU)",
            "Search Performance: Under 2s for 1000+ patients",
            "Transcription Latency: 2-3s real-time processing"
        ]),
        ("Security Validation", [
            "HIPAA Compliance: All requirements verified",
            "SQL Injection Prevention: Comprehensive testing",
            "Authentication Security: Token tampering protection",
            "Data Isolation: Cross-therapist access prevention",
            "Encryption: At-rest and in-transit validation"
        ]),
        ("Reliability Metrics", [
            "Uptime Target: 99.9% availability",
            "Error Rate: <1% for critical operations",
            "Data Integrity: 100% transaction consistency",
            "Backup Recovery: <1 hour RTO/RPO",
            "Graceful Degradation: Fallback mechanisms tested"
        ])
    ]
    
    for section, metrics in quality_sections:
        create_heading(doc, section, 3)
        for metric in metrics:
            doc.add_paragraph(metric, style='List Bullet')
    
    # 7. Deployment Requirements
    create_heading(doc, "7. DEPLOYMENT REQUIREMENTS", 1)
    
    create_heading(doc, "7.1 System Requirements", 2)
    
    create_heading(doc, "7.1.1 Hardware Requirements", 3)
    
    hw_configs = [
        ("Minimum Configuration", [
            "CPU: 4 cores, 2.5GHz",
            "RAM: 8GB",
            "Storage: 50GB SSD",
            "Network: 100Mbps local network",
            "OS: Ubuntu 20.04+ or equivalent"
        ]),
        ("Recommended Configuration", [
            "CPU: 8 cores, 3.0GHz",
            "RAM: 16GB",
            "Storage: 100GB NVMe SSD",
            "GPU: NVIDIA GTX 1660+ (optional, for faster AI)",
            "Network: Gigabit Ethernet",
            "OS: Ubuntu 22.04 LTS"
        ]),
        ("Production Configuration", [
            "CPU: 16 cores, 3.5GHz",
            "RAM: 32GB",
            "Storage: 500GB NVMe SSD + backup storage",
            "GPU: NVIDIA RTX 3070+ (recommended)",
            "Network: Redundant gigabit connections",
            "OS: Ubuntu 22.04 LTS with security hardening"
        ])
    ]
    
    for config, specs in hw_configs:
        create_heading(doc, config, 4)
        for spec in specs:
            doc.add_paragraph(spec, style='List Bullet')
    
    create_heading(doc, "7.2 Docker Containerization", 2)
    
    doc.add_paragraph("Container Structure:")
    container_services = [
        "backend-api: FastAPI application server",
        "backend-ws: WebSocket transcription server",
        "database: SQLite with volume persistence", 
        "models: AI model storage and serving"
    ]
    
    for service in container_services:
        doc.add_paragraph(service, style='List Bullet')
    
    create_heading(doc, "7.3 Security Configuration", 2)
    
    security_configs = [
        ("Network Security", [
            "Firewall: Restrict access to necessary ports only",
            "VPN: Optional VPN access for remote administration",
            "SSL/TLS: HTTPS certificates for production",
            "Network Isolation: Separate network segments for components"
        ]),
        ("Application Security", [
            "Environment Variables: Secure configuration management",
            "Secret Management: JWT keys and encryption keys",
            "User Permissions: Non-root container execution",
            "File Permissions: Restricted access to sensitive files"
        ]),
        ("HIPAA Compliance", [
            "Encryption: AES-256 for data at rest",
            "Access Logs: Comprehensive audit trails",
            "Backup Encryption: Encrypted backup storage",
            "User Training: Security awareness and procedures"
        ])
    ]
    
    for config, items in security_configs:
        create_heading(doc, config, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)